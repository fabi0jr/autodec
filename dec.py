import os
import re
import time
import PyPDF2
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
import shutil
from selenium.webdriver.support.ui import Select
from docx import Document
from docx.shared import Pt
import pyautogui
import pandas as pd
# Configurar o navegado
driver = webdriver.Chrome()

#Realizar login no site
driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
usuario = driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
senha = driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
botao_login = driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login

usuario.send_keys('03483401253')
senha.send_keys('carteira23')
botao_login.click()



planilha = pd.read_excel('cpfs.xlsx')

for index, row in planilha.iterrows():
    cpf = row['CPF']  # Pegar o CPF da linha atual

    def extrair_nome_municipio(endereco):
        # Lista de municípios
        mapeamento_municipios = {
            '413': 'ALVARAES',
            '414': 'AMATURA',
            '415':'Anamã',
            '416':'Anori',          
            '417':'Apuí',
            '418':'Atalaia do Norte',
            '419':'Autazes',
            '420':'Barcelos',
            '421':'Barreirinha',
            '422':'Benjamin Constant',
            '423':'Beruri',
            '424':'Boa Vista do Ramos',
            '425':'Boca do Acre',
            '426':'Borba',
            '427':'Caapiranga',
            '428':'Canutama',
            '429':'Carauari',
            '430':'Careiro',
            '431':'Careiro da Várzea',
            '432':'Coari',
            '433':'Codajás',
            '434':'Eirunepé',
            '435':'Envira',
            '436':'Fonte Boa',
            '437':'Guajará',
            '438':'Humaitá',
            '439':'Ipixuna',
            '440':'Iranduba',
            '441':'Itacoatiara',
            '442':'Itamarati',
            '443':'Itapiranga',
            '444':'Japurá',
            '445':'Juruá',
            '446':'Jutaí',
            '448':'Lábrea',
            '449':'Manacapuru',
            '450':'Manaquiri',
            '451':'Manaus',
            '452':'Manicoré',
            '453':'Maraã',
            '454':'Maués',
            '455':'Nhamundá',
            '456':'Nova Olinda do Norte',
            '457':'Novo Airão',
            '458':'Novo Aripuanã',
            '459':'Parintins',
            '460':'Pauini',
            '461':'Presidente Figueiredo',
            '462':'Rio Preto da Eva',
            '463':'Santa Isabel do Rio Negro',
            '464':'Santo Antônio do Içá',
            '465':'São Gabriel da Cachoeira',
            '466':'São Paulo de Olivença',
            '467':'São Sebastião do Uatumã',
            '468':'Silves',
            '469':'Tabatinga',
            '470':'Tapauá',
            '471':'Tefé',
            '472':'Tonantins',
            '473':'Uarini',
            '474':'Urucará',
            '475':'Urucurituba',
        }
        
        # Use expressões regulares para encontrar o nome do município no endereço
        for codigo, nome in mapeamento_municipios.items():
            if re.search(r'\b' + re.escape(nome) + r'\b', endereco, re.IGNORECASE):
                return nome
        
        return municipio  # Retorna isso se o município não for encontrado


    while True:
        # Selecione o elemento usando o XPath
        elemento = driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')

        # Faça algo com o elemento, por exemplo, clique nele
        elemento.click()
        
        # Localize a lista suspensa usando o XPath fornecido (a <div> que representa o menu)
        dropdown_div = driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')
        # Clique na <div> para expandir o menu
        dropdown_div.click()
        # Suponha que você tenha o CPF que deseja colar em uma variável chamada "cpf"
        
        caixacpf = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')  # Substitua pelo XPath do campo de entrada
        caixacpf.click()  # Clique no campo de entrada para garantir que ele está ativo
        caixacpf.send_keys(cpf)  # Cole o conteúdo da área de transferência (ou 'cmd' em vez de 'ctrl' no Mac)
        consultar = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')
        
        consultar.click()
        haProcesso = input("Aperte S se não houver cadastro:")
        if haProcesso.lower() == 's':
            # Atualize a coluna 'Processado' para marcar a linha como processada
            planilha.at[index, 'Não há processo'] = 'x'
            # Salve a planilha atualizada no mesmo arquivo Excel
            planilha.to_excel('cpfs.xlsx', index=False)
            break  # Saia do loop se o usuário não quiser continuar

        abadeclaração = driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr[1]/td[8]/a[2]')
        abadeclaração.click()
        #################################################################################
        nome_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cceaPessoaFisica_pfNome"]')
        # Use JavaScript para obter o valor do atributo 'value' do elemento
        nome_da_pagina = driver.execute_script("return arguments[0].value;", nome_element)

        cpf_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
        cpf_da_pagina = driver.execute_script("return arguments[0]. value;", cpf_element)

        propiedade_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmPropriedade"]')
        propriedade_da_pagina = driver.execute_script("return arguments[0]. value;", propiedade_element)

        endereco_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_txEnderecoPropriedade"]')
        endereco_da_pagina = driver.execute_script("return arguments[0]. value;", endereco_element)

        unloc_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_sgDistritoIdam"]')
        unloc_da_pagina = driver.execute_script("return arguments[0]. value;", unloc_element)

        latitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLatitude"]')
        latitude_da_pagina = driver.execute_script("return arguments[0]. value;", latitude_element)

        longitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLongitude"]')
        longitude_da_pagina = driver.execute_script("return arguments[0]. value;", longitude_element)

        muni_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_idMunicipio"]')
        muni_da_pagina = driver.execute_script("return arguments[0]. value;", muni_element)
        # Mapeamento de valores numéricos para nomes de municípios
        mapeamento_municipios = {
            '413': 'ALVARAES',
            '414': 'AMATURA',
            '415':'Anamã',
            '416':'Anori',          
            '417':'Apuí',
            '418':'Atalaia do Norte',
            '419':'Autazes',
            '420':'Barcelos',
            '421':'Barreirinha',
            '422':'Benjamin Constant',
            '423':'Beruri',
            '424':'Boa Vista do Ramos',
            '425':'Boca do Acre',
            '426':'Borba',
            '427':'Caapiranga',
            '428':'Canutama',
            '429':'Carauari',
            '430':'Careiro',
            '431':'Careiro da Várzea',
            '432':'Coari',
            '433':'Codajás',
            '434':'Eirunepé',
            '435':'Envira',
            '436':'Fonte Boa',
            '437':'Guajará',
            '438':'Humaitá',
            '439':'Ipixuna',
            '440':'Iranduba',
            '441':'Itacoatiara',
            '442':'Itamarati',
            '443':'Itapiranga',
            '444':'Japurá',
            '445':'Juruá',
            '446':'Jutaí',
            '448':'Lábrea',
            '449':'Manacapuru',
            '450':'Manaquiri',
            '451':'Manaus',
            '452':'Manicoré',
            '453':'Maraã',
            '454':'Maués',
            '455':'Nhamundá',
            '456':'Nova Olinda do Norte',
            '457':'Novo Airão',
            '458':'Novo Aripuanã',
            '459':'Parintins',
            '460':'Pauini',
            '461':'Presidente Figueiredo',
            '462':'Rio Preto da Eva',
            '463':'Santa Isabel do Rio Negro',
            '464':'Santo Antônio do Içá',
            '465':'São Gabriel da Cachoeira',
            '466':'São Paulo de Olivença',
            '467':'São Sebastião do Uatumã',
            '468':'Silves',
            '469':'Tabatinga',
            '470':'Tapauá',
            '471':'Tefé',
            '472':'Tonantins',
            '473':'Uarini',
            '474':'Urucará',
            '475':'Urucurituba',
        }

        # Suponha que você já tenha obtido o valor numérico do município
        valor_numerico = muni_da_pagina

        # Obtém o nome do município com base no valor numérico usando o mapeamento
        nome_municipio = mapeamento_municipios.get(valor_numerico, 'MUNICIPIO_DESCONHECIDO')

        atv1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaePrincipal"]')
        atv1_da_pagina = driver.execute_script("return arguments[0]. value;", atv1_element)

        qnt1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_qtdCnaePrincipalFormatado"]')
        qnt1_da_pagina = driver.execute_script("return arguments[0]. value;", qnt1_element)

        atv2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaeSecundario"]')
        atv2_da_pagina = driver.execute_script("return arguments[0]. value;", atv2_element)

        qnt2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_qtdCnaeSecundarioFormatado"]')
        qnt2_da_pagina = driver.execute_script("return arguments[0]. value;", qnt2_element)

        inicioatv_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_anoInicioAtividade"]')
        inicioatv_da_pagina = driver.execute_script("return arguments[0]. value;", inicioatv_element)

        numcontrole_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nrDeclaracaoUnidLocal"]')
        numcontrole_da_pagina = driver.execute_script("return arguments[0]. value;", numcontrole_element)

        ############################################################################
    # Selecione o elemento usando o XPath
        elemento = driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC1006"]')

        # Faça algo com o elemento, por exemplo, clique nele
        elemento.click()
        
        # Localize a lista suspensa usando o XPath fornecido (a <div> que representa o menu)
        dropdown_div = driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')

        # Clique na <div> para expandir o menu
        dropdown_div.click()
        caixarg = driver.find_element(By.XPATH,'//*[@id="formCceaPessoaFisica_pessoaFisicaTOA_pessoaFisica_cpfFormatado"]')
        caixarg.click()
        caixarg.send_keys(cpf)
        consultarrg = driver.find_element(By.XPATH,'//*[@id="formCceaPessoaFisica_pessoaFisica!pesquisarPessoaFisica"]')
        consultarrg.click()
        opcaorg = driver.find_element(By.XPATH,'//*[@id="tbPessoaFisica"]/tbody/tr/td[6]')
        opcaorg.click()


        rg_element = driver.find_element(By.XPATH, '//*[@id="formPessoaFisica_pessoaFisicaTOA_pessoaFisica_pfNumeroRg"]')
        rg_da_pagina = driver.execute_script("return arguments[0]. value;", rg_element)

        # Adicione a verificação condicional
        if not qnt1_da_pagina:
            qnt1_da_pagina = "0,5-HECTARES"

        if not qnt2_da_pagina:
            qnt2_da_pagina = "0,5-HECTARES"
        
        if not numcontrole_da_pagina:
            numcontrole_da_pagina = "00"
        # Use input() para obter o CPF do usuário

        ##################################################################################


        cpf = cpf_da_pagina
        rg = rg_da_pagina
        nome = nome_da_pagina
        propriedade = propriedade_da_pagina
        endereco = endereco_da_pagina
        unloc = unloc_da_pagina
        latitude = latitude_da_pagina
        longitude = longitude_da_pagina
        municipio = nome_municipio  # Nome do município com base no mapeamento
        atv1 = atv1_da_pagina
        qnt1 = qnt1_da_pagina
        atv2 = atv2_da_pagina
        qnt2 = qnt2_da_pagina
        inicioatv = inicioatv_da_pagina
        numcontrole = numcontrole_da_pagina

        # Crie um novo documento Word com base no modelo
    # Verifique a presença de atividades secundárias (atv2) e latitude
        if not atv2:
            if not latitude:
                doc = Document(r"I:\ARQUIVO DIGITAL CPR\fabio jr\autodec\modelo_declaracao_semcoordesematv2.docx.docx")
            else:
                doc = Document(r"I:\ARQUIVO DIGITAL CPR\fabio jr\autodec\modelo_declaracao_sematv2.docx.docx")
        elif not latitude:
            doc = Document(r"I:\ARQUIVO DIGITAL CPR\fabio jr\autodec\modelo_declaracao_semcoord.docx.docx")
        else:
            doc = Document(r"I:\ARQUIVO DIGITAL CPR\fabio jr\autodec\modelo_declaracao_completo.docx.docx")


        # Substitua as informações nas áreas específicas do documento
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if '(municipio)' in text:
                nome_municipio = extrair_nome_municipio(endereco)
                text = text.replace('(municipio)', nome_municipio)
            if '(unloc)' in text:
                text = text.replace('(unloc)', unloc)
            if '(numcontrole)' in text:
                text = text.replace('(numcontrole)', numcontrole)
            if '(nome)' in text:
                text = text.replace('(nome)', nome)
            if '(cpf)' in text:
                text = text.replace('(cpf)', cpf)
            if '(propriedade)' in text:
                text = text.replace('(propriedade)', propriedade)
            if '(endereço)' in text:
                text = text.replace('(endereço)', endereco)
            if '(atv1)' in text:
                text = text.replace('(atv1)', atv1)
            if '(qnt1)' in text:
                text = text.replace('(qnt1)', qnt1)
            if '(atv2)' in text:
                text = text.replace('(atv2)', atv2)
            if '(qnt2)' in text:
                text = text.replace('(qnt2)', qnt2)
            if '(inicioatv)' in text:
                text = text.replace('(inicioatv)', inicioatv)
            if '(latitude)' in text:
                text = text.replace('(latitude)', latitude)
            if '(longitude)' in text:
                text = text.replace('(longitude)', longitude)
            if '(rg)' in text:
                text = text.replace('(rg)', rg)
            #if '(ano)' in text:
                #text = text.replace('(ano)', ano)
            
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.size = Pt(12)  # Defina o tamanho da fonte conforme necessário

        # Salve o documento com um nome específico
        nome_do_arquivo = nome  # Solicitar um nome para o arquivo
        nome_do_arquivo = nome_do_arquivo + ".docx"  # Adicione a extensão .docx
        doc.save(nome_do_arquivo)  # Salve o documento com o nome fornecido

        break

        
# Feche o navegador
driver.quit()




